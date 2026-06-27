from __future__ import annotations

import csv
import importlib
import io
import json
import logging
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
import uuid
import zipfile
from dataclasses import dataclass, field
from html import unescape
from pathlib import Path
from typing import TYPE_CHECKING, Any, Callable

from app.config import settings
from app.schemas.callback import CallbackResult
from app.schemas.task_message import DocumentProcessMessage

if TYPE_CHECKING:
    from app.infrastructure.storage.minio_client import MinioDocumentClient

logger = logging.getLogger(__name__)

UNRECOGNIZED_STRUCTURE = "未识别结构"
UNRECOGNIZED_DIRECTORY = "未识别目录"
UNASSIGNED_BLOCK = "未归属解析块"


@dataclass
class ParsedBlock:
    """解析后的解析块，解析块是分片的来源单元。"""

    block_id: str
    block_name: str
    block_type: str
    title_path: str
    text: str
    parse_block_id: str | None = None
    parse_block_name: str | None = None
    parse_block_page_start: int | None = None
    parse_block_page_end: int | None = None
    parse_block_metadata: dict[str, Any] = field(default_factory=dict)
    page_start: int | None = None
    page_end: int | None = None
    sheet_name: str | None = None
    row_start: int | None = None
    row_end: int | None = None


@dataclass
class ParsedSection:
    """文档目录结构节点。"""

    section_no: int
    section_id: str
    parent_section_id: str | None
    title: str
    title_path: str
    level: int
    blocks: list[ParsedBlock] = field(default_factory=list)


@dataclass
class PdfElement:
    """PDF 页面元素，先从页面版式抽取，再转换为目录结构和解析块。"""

    page_no: int
    element_type: str
    text: str
    bbox: tuple[float, float, float, float]
    font_size: float = 0.0
    bold: bool = False
    source: str = "native"
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class PdfParseBlock:
    """PDF 按配置页码范围拆出的解析块。"""

    block_id: str
    block_name: str
    page_start: int
    page_end: int
    elements: list[PdfElement]
    title_candidates: list[str] = field(default_factory=list)
    element_counts: dict[str, int] = field(default_factory=dict)
    ocr_status: str = "not_required"
    ocr_message: str | None = None
    local_directory_tree: list[dict[str, Any]] = field(default_factory=list)


@dataclass
class Chunk:
    """知识分片结果。"""

    chunk_id: str
    chunk_no: int
    chunk_text: str
    section_id: str | None
    title_path: str
    block_ids: list[str]
    block_names: list[str]
    parse_block_id: str | None
    parse_block_name: str | None
    token_count: int
    char_count: int
    page_start: int | None = None
    page_end: int | None = None
    sheet_name: str | None = None
    row_start: int | None = None
    row_end: int | None = None


class ParseChunkPipeline:
    """文档解析与分片应用服务，负责生成结构、解析块和分片产物。"""

    def __init__(self, storage_client: "MinioDocumentClient") -> None:
        self._storage_client = storage_client

    def run(
        self,
        message: DocumentProcessMessage,
        progress_callback: Callable[[int], None] | None = None,
    ) -> CallbackResult:
        logger.info(
            "解析分片流水线开始 task_id=%s document_id=%s version_id=%s object_key=%s",
            message.taskId,
            message.documentId,
            message.versionId,
            message.objectKey,
        )
        file_bytes = self._storage_client.download_bytes(message.bucketName, message.objectKey)
        self._report_progress(progress_callback, 20)
        parser_name = self._parser_name(message)
        sections = self._parse_document(file_bytes, message)
        self._report_progress(progress_callback, 68)
        chunks = self._split_chunks(sections, message)
        self._report_progress(progress_callback, 78)
        base_key = f"parsed/{message.knowledgeBaseId}/{message.documentId}/{message.versionId}/{message.taskId}"

        sections_key = f"{base_key}/sections.jsonl"
        structure_key = f"{base_key}/document_structure.json"
        chunks_key = f"{base_key}/chunks.jsonl"

        self._storage_client.upload_text(
            message.bucketName,
            sections_key,
            self._build_sections_jsonl(sections),
            "application/x-jsonlines",
        )
        self._report_progress(progress_callback, 85)
        self._storage_client.upload_text(
            message.bucketName,
            structure_key,
            self._build_structure_json(message, sections, parser_name),
            "application/json",
        )
        self._report_progress(progress_callback, 92)
        self._storage_client.upload_text(
            message.bucketName,
            chunks_key,
            self._build_chunks_jsonl(message, chunks),
            "application/x-jsonlines",
        )
        self._report_progress(progress_callback, 97)

        token_count = sum(chunk.token_count for chunk in chunks)
        logger.info(
            "解析分片流水线完成 task_id=%s section_count=%s block_count=%s chunk_count=%s token_count=%s",
            message.taskId,
            len(sections),
            sum(len(section.blocks) for section in sections),
            len(chunks),
            token_count,
        )
        return CallbackResult(
            sectionsObjectKey=sections_key,
            structureObjectKey=structure_key,
            chunksObjectKey=chunks_key,
            chunkCount=len(chunks),
            tokenCount=token_count,
            parserName=parser_name,
        )

    def _report_progress(self, callback: Callable[[int], None] | None, progress: int) -> None:
        if callback is None:
            return
        try:
            callback(progress)
        except Exception as exception:
            # 中间进度回调失败不能丢弃已完成的解析结果，最终成功/失败回调仍由消费者负责。
            logger.warning("文档处理中间进度回调失败 progress=%s error=%s", progress, exception)

    def _parse_document(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[ParsedSection]:
        file_type = message.fileType.lower()
        if file_type in {"md", "markdown"}:
            sections = self._parse_markdown(self._decode_text(file_bytes), message)
            return self._ensure_parse_block_metadata(sections, message)
        if file_type == "csv":
            sections = self._parse_csv(file_bytes, message)
            return self._ensure_parse_block_metadata(sections, message)
        if file_type in {"xlsx", "xls"}:
            sections = self._parse_excel(file_bytes, message)
            return self._ensure_parse_block_metadata(sections, message)
        if file_type in {"docx", "doc"}:
            sections = self._parse_docx(file_bytes, message)
            return self._ensure_parse_block_metadata(sections, message)
        if file_type == "pdf":
            sections = self._parse_pdf(file_bytes, message)
            return self._ensure_parse_block_metadata(sections, message)
        sections = self._parse_plain_text(self._decode_text(file_bytes), message)
        return self._ensure_parse_block_metadata(sections, message)

    def _parser_name(self, message: DocumentProcessMessage) -> str:
        if message.fileType.lower() == "pdf":
            return self._pdf_parser_type(message)
        return "rule-structure-parser"

    def _pdf_parser_type(self, message: DocumentProcessMessage) -> str:
        """从版本分片配置快照读取 PDF 解析器；历史消息没有配置时默认 Docling。"""
        parser_type = "docling"
        if message.chunkConfigSnapshot:
            try:
                snapshot = json.loads(message.chunkConfigSnapshot)
                parser_type = str(snapshot.get("parserType") or snapshot.get("parser_type") or parser_type).lower()
            except (TypeError, ValueError, json.JSONDecodeError) as exception:
                logger.warning(
                    "PDF 分片配置快照解析失败，使用默认 Docling task_id=%s document_id=%s stage=%s error=%s",
                    message.taskId,
                    message.documentId,
                    message.stageCode,
                    exception,
                )
        if parser_type not in {"docling", "mineru", "pdf"}:
            raise ValueError(f"不支持的 PDF 解析器: {parser_type}")
        return parser_type

    def _parse_markdown(self, text: str, message: DocumentProcessMessage) -> list[ParsedSection]:
        lines = self._normalize_text(text).splitlines()
        sections: list[ParsedSection] = []
        heading_stack: list[tuple[int, str]] = []
        current_section: ParsedSection | None = None
        buffer: list[str] = []
        table_buffer: list[str] = []
        code_buffer: list[str] = []
        in_code = False
        section_no = 0

        def ensure_section() -> ParsedSection:
            nonlocal current_section, section_no
            if current_section is None:
                title_path = self._current_title_path(heading_stack)
                section_no += 1
                current_section = ParsedSection(
                    section_no=section_no,
                    section_id=self._uuid32(),
                    parent_section_id=None,
                    title=title_path.split(" / ")[-1],
                    title_path=title_path,
                    level=max(1, title_path.count(" / ") + 1),
                    blocks=[],
                )
                sections.append(current_section)
            return current_section

        def flush_text() -> None:
            nonlocal buffer
            content = "\n".join(buffer).strip()
            if not content:
                buffer = []
                return
            section = ensure_section()
            section.blocks.append(
                ParsedBlock(
                    block_id=self._uuid32(),
                    block_name=f"{section.title_path} · 正文",
                    block_type="paragraph",
                    title_path=section.title_path,
                    text=content,
                )
            )
            buffer = []

        def flush_table() -> None:
            nonlocal table_buffer
            if not table_buffer:
                return
            section = ensure_section()
            section.blocks.append(
                ParsedBlock(
                    block_id=self._uuid32(),
                    block_name=f"{section.title_path} · Markdown 表格",
                    block_type="table",
                    title_path=section.title_path,
                    text="\n".join(table_buffer).strip(),
                )
            )
            table_buffer = []

        def flush_code() -> None:
            nonlocal code_buffer
            if not code_buffer:
                return
            section = ensure_section()
            section.blocks.append(
                ParsedBlock(
                    block_id=self._uuid32(),
                    block_name=f"{section.title_path} · 代码块",
                    block_type="code",
                    title_path=section.title_path,
                    text="\n".join(code_buffer).strip(),
                )
            )
            code_buffer = []

        for raw_line in lines:
            line = raw_line.rstrip()
            if line.strip().startswith("```"):
                if in_code:
                    flush_code()
                    in_code = False
                else:
                    flush_text()
                    flush_table()
                    in_code = True
                continue
            if in_code:
                code_buffer.append(line)
                continue
            heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
            if heading:
                flush_text()
                flush_table()
                level = len(heading.group(1))
                title = self._clean_heading(heading.group(2))
                heading_stack = [item for item in heading_stack if item[0] < level]
                heading_stack.append((level, title))
                title_path = self._current_title_path(heading_stack)
                section_no += 1
                current_section = ParsedSection(
                    section_no=section_no,
                    section_id=self._uuid32(),
                    parent_section_id=None,
                    title=title,
                    title_path=title_path,
                    level=level,
                    blocks=[
                        ParsedBlock(
                            block_id=self._uuid32(),
                            block_name=f"{title_path} · 标题",
                            block_type="title",
                            title_path=title_path,
                            text=title,
                        )
                    ],
                )
                sections.append(current_section)
                continue
            if self._is_markdown_table_line(line):
                flush_text()
                table_buffer.append(line)
                continue
            if table_buffer and not self._is_markdown_table_line(line):
                flush_table()
            buffer.append(line)
        flush_text()
        flush_table()
        flush_code()
        return sections or [self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))]

    def _parse_plain_text(self, text: str, message: DocumentProcessMessage) -> list[ParsedSection]:
        lines = self._normalize_text(text).splitlines()
        sections: list[ParsedSection] = []
        current_title = UNRECOGNIZED_STRUCTURE
        buffer: list[str] = []
        section_no = 0

        def flush() -> None:
            nonlocal buffer, section_no
            content = "\n".join(buffer).strip()
            if not content:
                buffer = []
                return
            section_no += 1
            sections.append(self._section_from_text(section_no, current_title, content))
            buffer = []

        for raw_line in lines:
            line = raw_line.strip()
            if self._looks_like_heading(line):
                flush()
                current_title = self._clean_heading(line)
                continue
            buffer.append(raw_line)
        flush()
        return sections or [self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))]

    def _parse_csv(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[ParsedSection]:
        text = self._decode_text(file_bytes)
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            return [self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))]
        header = rows[0]
        data_rows = rows[1:] or []
        return [self._table_section(1, "CSV 表格", "CSV", header, data_rows, 2)]

    def _parse_excel(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[ParsedSection]:
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning("openpyxl 未安装，按普通文本兜底解析 task_id=%s", message.taskId)
            return self._parse_plain_text(self._decode_text(file_bytes), message)

        try:
            workbook = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            sections: list[ParsedSection] = []
            for index, sheet in enumerate(workbook.worksheets, start=1):
                values = [[self._cell_text(cell) for cell in row] for row in sheet.iter_rows(values_only=True)]
                values = [row for row in values if any(cell for cell in row)]
                if not values:
                    continue
                header = values[0]
                data_rows = values[1:]
                sections.append(self._table_section(index, f"工作表：{sheet.title}", sheet.title, header, data_rows, 2))
            return sections or [self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))]
        except Exception as exception:
            logger.warning("Excel 解析失败 task_id=%s error=%s", message.taskId, exception)
            return [self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))]

    def _parse_docx(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[ParsedSection]:
        try:
            from docx import Document

            document = Document(io.BytesIO(file_bytes))
            sections: list[ParsedSection] = []
            current_section: ParsedSection | None = None
            title_stack: list[tuple[int, str]] = []

            def ensure_section() -> ParsedSection:
                nonlocal current_section
                if current_section is None:
                    current_section = ParsedSection(
                        section_no=len(sections) + 1,
                        section_id=self._uuid32(),
                        parent_section_id=None,
                        title=UNRECOGNIZED_STRUCTURE,
                        title_path=UNRECOGNIZED_STRUCTURE,
                        level=1,
                        blocks=[],
                    )
                    sections.append(current_section)
                return current_section

            for paragraph in document.paragraphs:
                text = self._normalize_text(paragraph.text)
                if not text:
                    continue
                style_name = paragraph.style.name if paragraph.style is not None else ""
                heading_level = self._docx_heading_level(style_name, text)
                if heading_level:
                    title = self._clean_heading(text)
                    title_stack = [item for item in title_stack if item[0] < heading_level]
                    title_stack.append((heading_level, title))
                    title_path = self._current_title_path(title_stack)
                    current_section = ParsedSection(
                        section_no=len(sections) + 1,
                        section_id=self._uuid32(),
                        parent_section_id=None,
                        title=title,
                        title_path=title_path,
                        level=heading_level,
                        blocks=[
                            ParsedBlock(
                                block_id=self._uuid32(),
                                block_name=f"{title_path} · 标题",
                                block_type="title",
                                title_path=title_path,
                                text=title,
                            )
                        ],
                    )
                    sections.append(current_section)
                    continue
                section = ensure_section()
                section.blocks.append(
                    ParsedBlock(
                        block_id=self._uuid32(),
                        block_name=f"{section.title_path} · 段落",
                        block_type="paragraph",
                        title_path=section.title_path,
                        text=text,
                    )
                )

            for table_index, table in enumerate(document.tables, start=1):
                section = ensure_section()
                rows = [[self._normalize_text(cell.text) for cell in row.cells] for row in table.rows]
                header = rows[0] if rows else []
                body_rows = rows[1:] if len(rows) > 1 else []
                section.blocks.append(
                    ParsedBlock(
                        block_id=self._uuid32(),
                        block_name=f"{section.title_path} · Word 表格 {table_index}",
                        block_type="table",
                        title_path=section.title_path,
                        text=self._table_text(header, body_rows),
                    )
                )

            image_count = len(document.inline_shapes)
            if image_count:
                section = ensure_section()
                for image_index in range(1, image_count + 1):
                    section.blocks.append(
                        ParsedBlock(
                            block_id=self._uuid32(),
                            block_name=f"{section.title_path} · 图片 {image_index}",
                            block_type="image",
                            title_path=section.title_path,
                            text=f"[图片] Word 文档图片 {image_index}",
                        )
                    )
            return [section for section in sections if section.blocks] or [
                self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))
            ]
        except ImportError:
            logger.warning("python-docx 未安装，使用 docx XML 兜底解析 task_id=%s", message.taskId)
        except Exception as exception:
            logger.warning("python-docx 解析失败 task_id=%s error=%s", message.taskId, exception)
        return self._parse_docx_xml(file_bytes, message)

    def _parse_docx_xml(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[ParsedSection]:
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
                xml_text = archive.read("word/document.xml").decode("utf-8", errors="ignore")
            plain_text = " ".join(unescape(value) for value in re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml_text))
            return self._parse_plain_text(plain_text, message)
        except Exception as exception:
            logger.warning("docx XML 兜底解析失败 task_id=%s error=%s", message.taskId, exception)
            return [self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))]

    def _parse_pdf(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[ParsedSection]:
        parser_type = self._pdf_parser_type(message)
        logger.info(
            "PDF 解析开始 task_id=%s document_id=%s stage=%s parser=%s",
            message.taskId,
            message.documentId,
            message.stageCode,
            parser_type,
        )
        if parser_type == "docling":
            parse_blocks = self._parse_pdf_with_docling(file_bytes, message)
        elif parser_type == "mineru":
            parse_blocks = self._parse_pdf_with_mineru(file_bytes, message)
        else:
            parse_blocks = self._parse_pdf_with_native_layout(file_bytes, message)
        if not parse_blocks:
            raise ValueError(f"{parser_type} 未生成 PDF 解析块")
        sections = self._sections_from_pdf_parse_blocks(parse_blocks)
        if not sections:
            raise ValueError(f"{parser_type} 未生成 PDF 目录和正文结构")
        logger.info(
            "PDF 解析完成 task_id=%s document_id=%s stage=%s parser=%s parse_block_count=%s section_count=%s",
            message.taskId,
            message.documentId,
            message.stageCode,
            parser_type,
            len(parse_blocks),
            len(sections),
        )
        return sections

    def _parse_pdf_with_native_layout(
        self, file_bytes: bytes, message: DocumentProcessMessage
    ) -> list[PdfParseBlock]:
        try:
            import fitz
        except ImportError as exception:
            raise RuntimeError("PDF 原生解析器需要安装 PyMuPDF") from exception

        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as document:
                return self._build_pdf_parse_blocks(file_bytes, document, enable_mineru_ocr=False)
        except Exception as exception:
            logger.exception("PDF 原生版面解析失败 task_id=%s document_id=%s stage=%s", message.taskId, message.documentId, message.stageCode)
            raise RuntimeError(f"PDF 原生版面解析失败: {exception}") from exception

    def _parse_pdf_with_docling(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[PdfParseBlock]:
        try:
            from docling.document_converter import DocumentConverter
        except ImportError as exception:
            raise RuntimeError(f"未安装 Docling，请先执行: {sys.executable} -m pip install 'docling>=2.0.0'") from exception

        try:
            with tempfile.TemporaryDirectory(prefix="aistudio-docling-") as temp_dir:
                pdf_path = Path(temp_dir) / self._safe_pdf_file_name(message.fileName)
                pdf_path.write_bytes(file_bytes)
                document = DocumentConverter().convert(pdf_path).document
                elements = self._docling_elements(document)
            return self._build_external_pdf_parse_blocks(elements, self._pdf_page_ranges(file_bytes), "docling")
        except Exception as exception:
            logger.exception("Docling 解析失败 task_id=%s document_id=%s stage=%s", message.taskId, message.documentId, message.stageCode)
            raise RuntimeError(f"Docling 解析失败: {exception}") from exception

    def _parse_pdf_with_mineru(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[PdfParseBlock]:
        executable = self._resolve_mineru_command()
        if not executable:
            raise RuntimeError(
                "未找到 MinerU CLI，请先安装并确认 mineru 已加入 PATH: "
                f"{sys.executable} -m pip install 'mineru[all]>=2.0.0'"
            )
        try:
            with tempfile.TemporaryDirectory(prefix="aistudio-mineru-") as temp_dir:
                temp_path = Path(temp_dir)
                pdf_path = temp_path / self._safe_pdf_file_name(message.fileName)
                output_path = temp_path / "output"
                pdf_path.write_bytes(file_bytes)
                command = [executable, "-p", str(pdf_path), "-o", str(output_path)]
                completed = subprocess.run(
                    command,
                    capture_output=True,
                    text=True,
                    timeout=settings.mineru_timeout_seconds,
                    check=False,
                    encoding="utf-8",
                    errors="replace",
                )
                if completed.returncode != 0:
                    error = (completed.stderr or completed.stdout or "未知错误").strip()[-1000:]
                    raise RuntimeError(f"MinerU CLI exit_code={completed.returncode}: {error}")
                candidates = list(output_path.rglob("*_content_list_v2.json"))
                if not candidates:
                    candidates = list(output_path.rglob("*_content_list.json"))
                if not candidates:
                    raise RuntimeError("MinerU 未生成 content_list_v2.json 或 content_list.json")
                raw_pages = json.loads(candidates[0].read_text(encoding="utf-8"))
                elements = self._mineru_elements(raw_pages)
            return self._build_external_pdf_parse_blocks(elements, self._pdf_page_ranges(file_bytes), "mineru")
        except Exception as exception:
            logger.exception("MinerU 解析失败 task_id=%s document_id=%s stage=%s", message.taskId, message.documentId, message.stageCode)
            raise RuntimeError(f"MinerU 解析失败: {exception}") from exception

    def _safe_pdf_file_name(self, file_name: str) -> str:
        safe_name = Path(file_name or "document.pdf").name
        return safe_name if safe_name.lower().endswith(".pdf") else f"{safe_name}.pdf"

    def _pdf_page_ranges(self, file_bytes: bytes) -> list[tuple[int, int]]:
        try:
            from pypdf import PdfReader
        except ImportError as exception:
            raise RuntimeError("PDF 分页需要安装 pypdf") from exception
        page_count = len(PdfReader(io.BytesIO(file_bytes)).pages)
        if page_count <= 0:
            raise ValueError("PDF 不包含可解析页面")
        block_pages = max(1, settings.pdf_parse_block_pages)
        return [
            (page_start, min(page_count, page_start + block_pages - 1))
            for page_start in range(1, page_count + 1, block_pages)
        ]

    def _docling_elements(self, document: Any) -> list[PdfElement]:
        elements: list[PdfElement] = []
        for item, level in document.iterate_items(traverse_pictures=True):
            provenance = getattr(item, "prov", None) or []
            if not provenance:
                continue
            label_value = getattr(item, "label", "text")
            label = str(getattr(label_value, "value", label_value)).lower()
            text = self._docling_item_text(item, document, label)
            element_type = self._docling_element_type(label, text)
            if not text and element_type != "image":
                continue
            prov = provenance[0]
            page_no = max(1, int(getattr(prov, "page_no", 1)))
            metadata: dict[str, Any] = {"docling_label": label}
            if element_type == "title":
                # Docling 的 SectionHeaderItem.level 才是文档标题级别；iterate_items 返回的是遍历深度。
                item_level = getattr(item, "level", None)
                metadata["heading_level"] = (
                    1 if label == "title"
                    else max(2, min(int(item_level or level or 1) + 1, 6))
                )
            elements.append(
                PdfElement(
                    page_no=page_no,
                    element_type=element_type,
                    text=text,
                    bbox=self._docling_bbox(getattr(prov, "bbox", None), document, page_no),
                    font_size=max(10.0, 24.0 - (int(metadata.get("heading_level", 1)) - 1) * 2)
                    if element_type == "title" else 0.0,
                    bold=element_type == "title",
                    source="docling",
                    metadata=metadata,
                )
            )
        return sorted(elements, key=lambda item: (item.page_no, item.bbox[1], item.bbox[0]))

    def _docling_item_text(self, item: Any, document: Any, label: str) -> str:
        if label == "table":
            try:
                return item.export_to_dataframe(doc=document).to_csv(index=False).strip()
            except Exception as exception:
                logger.warning("Docling 表格转换失败 label=%s error=%s", label, exception)
                return str(getattr(item, "text", "") or "").strip()
        text = str(getattr(item, "text", "") or "").strip()
        if not text and label == "picture":
            try:
                return str(item.caption_text(document) or "").strip()
            except Exception:
                return ""
        return text

    def _docling_element_type(self, label: str, text: str) -> str:
        if label in {"title", "section_header"}:
            return "title"
        if label == "table":
            return "table"
        if label == "picture" and not text:
            return "image"
        if label in {"page_header", "page_footer"}:
            return "page_number"
        return "paragraph"

    def _docling_bbox(self, bbox: Any, document: Any, page_no: int) -> tuple[float, float, float, float]:
        result = self._bbox_values(bbox)
        origin_value = getattr(bbox, "coord_origin", "")
        origin = str(getattr(origin_value, "value", origin_value)).upper()
        if "BOTTOMLEFT" not in origin:
            return result
        page = getattr(document, "pages", {}).get(page_no)
        page_height = float(getattr(getattr(page, "size", None), "height", 0.0))
        if page_height <= 0:
            return result
        x0, y0, x1, y1 = result
        return (x0, page_height - y1, x1, page_height - y0)

    def _resolve_mineru_command(self) -> str | None:
        command_path = Path(settings.mineru_command)
        if command_path.is_absolute() and command_path.is_file():
            return str(command_path)
        return shutil.which(settings.mineru_command)

    def _mineru_elements(self, raw_pages: Any) -> list[PdfElement]:
        pages = raw_pages if isinstance(raw_pages, list) else raw_pages.get("pages", [])
        if not isinstance(pages, list):
            return []
        grouped = bool(pages) and all(isinstance(page, list) for page in pages)
        items = [
            (page_no, item)
            for page_no, page in enumerate(pages, start=1)
            for item in (page if isinstance(page, list) else [page])
        ]
        elements: list[PdfElement] = []
        for grouped_page_no, item in items:
            if not isinstance(item, dict):
                continue
            if grouped:
                page_no = grouped_page_no
            elif item.get("page_idx") is not None:
                page_no = int(item["page_idx"]) + 1
            else:
                page_no = int(item.get("page_no") or 1)
            source_type = str(item.get("type") or "paragraph").lower()
            content = item.get("content")
            text = self._mineru_content_text(source_type, content, item)
            if source_type in {"page_header", "page_footer", "header", "footer", "page_number", "page_aside_text"}:
                element_type = "page_number"
            elif source_type == "title":
                element_type = "title"
            elif source_type == "table":
                element_type = "table"
            elif source_type in {"image", "chart"} and not text:
                element_type = "image"
            else:
                element_type = "paragraph"
            if not text and element_type != "image":
                continue
            heading_level = content.get("level") if source_type == "title" and isinstance(content, dict) else item.get("text_level")
            metadata: dict[str, Any] = {"mineru_type": source_type}
            if heading_level is not None:
                metadata["heading_level"] = heading_level
            elements.append(
                PdfElement(
                    page_no=max(1, page_no),
                    element_type=element_type,
                    text=text,
                    bbox=self._bbox_values(item.get("bbox")),
                    font_size=max(10.0, 24.0 - (int(heading_level or 1) - 1) * 2)
                    if element_type == "title" else 0.0,
                    bold=element_type == "title",
                    source="mineru",
                    metadata=metadata,
                )
            )
        return sorted(elements, key=lambda item: (item.page_no, item.bbox[1], item.bbox[0]))

    def _mineru_content_text(self, source_type: str, content: Any, item: dict[str, Any]) -> str:
        field_by_type = {
            "title": "title_content", "paragraph": "paragraph_content", "text": "paragraph_content",
            "table": "table_content", "image": "image_caption", "chart": "chart_caption",
        }
        selected = content.get(field_by_type.get(source_type, "")) if isinstance(content, dict) else content
        if selected is None and isinstance(content, dict):
            selected = content.get("caption") or content
        if selected is None:
            selected = item.get("text") or item.get("table_body")
        return self._flatten_parser_text(selected).strip()

    def _flatten_parser_text(self, value: Any) -> str:
        if isinstance(value, str):
            return value
        if isinstance(value, list):
            return "".join(self._flatten_parser_text(item) for item in value)
        if isinstance(value, dict):
            for key in (
                "content", "text", "html", "markdown", "title_content", "paragraph_content",
                "table_content", "table_body", "table_caption", "image_caption", "chart_caption", "list_items",
            ):
                if key in value:
                    return self._flatten_parser_text(value[key])
        return ""

    def _bbox_values(self, bbox: Any) -> tuple[float, float, float, float]:
        if bbox is None:
            return (0.0, 0.0, 0.0, 0.0)
        if isinstance(bbox, (list, tuple)):
            values = [float(value) for value in bbox[:4]]
        else:
            values = [float(getattr(bbox, name, 0.0)) for name in ("l", "t", "r", "b")]
        if len(values) < 4:
            return (0.0, 0.0, 0.0, 0.0)
        x0, y0, x1, y1 = values
        return (min(x0, x1), min(y0, y1), max(x0, x1), max(y0, y1))

    def _build_external_pdf_parse_blocks(
        self, elements: list[PdfElement], page_ranges: list[tuple[int, int]], parser_type: str
    ) -> list[PdfParseBlock]:
        parse_blocks: list[PdfParseBlock] = []
        for block_index, (page_start, page_end) in enumerate(page_ranges, start=1):
            block_elements = [item for item in elements if page_start <= item.page_no <= page_end]
            title_summary = self._pdf_parse_block_title_summary(block_elements)
            block_name = (
                f"全文解析块｜第 {page_start}-{page_end} 页｜{title_summary}"
                if len(page_ranges) == 1 else
                f"解析块 {block_index:02d}｜第 {page_start}-{page_end} 页｜{title_summary}"
            )
            parse_blocks.append(
                PdfParseBlock(
                    block_id=self._uuid32(),
                    block_name=block_name,
                    page_start=page_start,
                    page_end=page_end,
                    elements=block_elements,
                    title_candidates=self._pdf_title_candidates(block_elements),
                    element_counts=self._pdf_element_counts(block_elements),
                    ocr_status=f"{parser_type}_success",
                )
            )
        return parse_blocks

    def _parse_pdf_text_fallback(self, file_bytes: bytes, message: DocumentProcessMessage) -> list[ParsedSection]:
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            sections: list[ParsedSection] = []
            for page_index, page in enumerate(reader.pages, start=1):
                text = self._normalize_text(page.extract_text() or "")
                if not text:
                    continue
                section = self._section_from_text(page_index, f"未识别结构 / 第 {page_index} 页", text, page_index, page_index)
                sections.append(section)
            return sections or [self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))]
        except Exception as exception:
            logger.warning("PDF 文本兜底解析失败 task_id=%s error=%s", message.taskId, exception)
            return [self._section_from_text(1, UNRECOGNIZED_STRUCTURE, self._fallback_text(message))]

    def _build_pdf_parse_blocks(
        self, file_bytes: bytes, document: Any, enable_mineru_ocr: bool = False
    ) -> list[PdfParseBlock]:
        page_count = len(document)
        if page_count <= 0:
            return []

        plumber_pdf = self._open_pdfplumber(file_bytes)
        block_pages = max(1, settings.pdf_parse_block_pages)
        page_ranges: list[tuple[int, int]]
        if page_count <= block_pages:
            page_ranges = [(1, page_count)]
        else:
            page_ranges = [
                (page_start, min(page_count, page_start + block_pages - 1))
                for page_start in range(1, page_count + 1, block_pages)
            ]

        parse_blocks: list[PdfParseBlock] = []
        try:
            for block_index, (page_start, page_end) in enumerate(page_ranges, start=1):
                elements = self._extract_pdf_elements(document, plumber_pdf, page_start, page_end)
                if enable_mineru_ocr:
                    mineru_status, mineru_message, mineru_elements = self._recognize_pdf_images_with_mineru(
                        file_bytes=file_bytes,
                        page_start=page_start,
                        page_end=page_end,
                        image_elements=[item for item in elements if item.element_type == "image"],
                    )
                else:
                    mineru_status, mineru_message, mineru_elements = "not_required", None, []
                elements.extend(mineru_elements)
                elements = sorted(elements, key=lambda item: (item.page_no, item.bbox[1], item.bbox[0]))
                title_summary = self._pdf_parse_block_title_summary(elements)
                title_candidates = self._pdf_title_candidates(elements)
                element_counts = self._pdf_element_counts(elements)
                if page_count <= block_pages:
                    block_name = f"全文解析块｜第 {page_start}-{page_end} 页｜{title_summary}"
                else:
                    block_name = f"解析块 {block_index:02d}｜第 {page_start}-{page_end} 页｜{title_summary}"
                parse_blocks.append(
                    PdfParseBlock(
                        block_id=self._uuid32(),
                        block_name=block_name,
                        page_start=page_start,
                        page_end=page_end,
                        elements=elements,
                        title_candidates=title_candidates,
                        element_counts=element_counts,
                        ocr_status=mineru_status,
                        ocr_message=mineru_message,
                    )
                )
        finally:
            if plumber_pdf is not None:
                plumber_pdf.close()
        return parse_blocks

    def _extract_pdf_elements(
        self,
        document: Any,
        plumber_pdf: Any | None,
        page_start: int,
        page_end: int,
    ) -> list[PdfElement]:
        raw_text_elements: list[PdfElement] = []
        table_elements: list[PdfElement] = []
        image_elements: list[PdfElement] = []
        for page_index in range(page_start, page_end + 1):
            fitz_page = document[page_index - 1]
            plumber_page = plumber_pdf.pages[page_index - 1] if plumber_pdf is not None else None
            page_table_elements, table_bboxes = self._extract_pdf_table_elements(plumber_page, fitz_page, page_index)
            table_elements.extend(page_table_elements)
            if plumber_page is not None:
                raw_text_elements.extend(self._extract_pdf_text_lines_from_plumber(plumber_page, page_index, table_bboxes))
            else:
                raw_text_elements.extend(self._extract_pdf_text_lines_from_fitz(fitz_page, page_index, table_bboxes))
            image_elements.extend(self._extract_pdf_image_elements(fitz_page, page_index))
        body_sizes = [item.font_size for item in raw_text_elements if item.font_size > 0]
        base_size = self._pdf_body_font_size(body_sizes)
        text_elements = [self._classify_pdf_text_element(item, base_size) for item in raw_text_elements]
        return sorted(text_elements + table_elements + image_elements, key=lambda item: (item.page_no, item.bbox[1], item.bbox[0]))

    def _recognize_pdf_images_with_mineru(
        self,
        file_bytes: bytes,
        page_start: int,
        page_end: int,
        image_elements: list[PdfElement],
    ) -> tuple[str, str | None, list[PdfElement]]:
        if not image_elements:
            return "not_required", None, []
        if not settings.enable_mineru_ocr:
            logger.info(
                "MinerU OCR 未启用，跳过图片识别 page_start=%s page_end=%s image_count=%s",
                page_start,
                page_end,
                len(image_elements),
            )
            return "skipped_disabled", "MinerU OCR 未启用", []
        if not settings.mineru_ocr_adapter:
            logger.warning(
                "MinerU OCR 已启用但未配置适配器，跳过图片识别 page_start=%s page_end=%s image_count=%s",
                page_start,
                page_end,
                len(image_elements),
            )
            return "skipped_no_adapter", "未配置 MinerU OCR 适配器", []

        try:
            adapter = self._load_ocr_adapter(settings.mineru_ocr_adapter)
            raw_result = adapter(
                file_bytes=file_bytes,
                page_start=page_start,
                page_end=page_end,
                image_elements=[
                    {
                        "page_no": item.page_no,
                        "bbox": item.bbox,
                        "text": item.text,
                        "metadata": item.metadata,
                    }
                    for item in image_elements
                ],
            )
            ocr_elements = self._normalize_ocr_elements(raw_result, page_start)
            return "success" if ocr_elements else "empty", None, ocr_elements
        except Exception as exception:
            logger.warning(
                "MinerU 图片识别失败 page_start=%s page_end=%s image_count=%s error=%s",
                page_start,
                page_end,
                len(image_elements),
                exception,
            )
            return "failed", str(exception)[:200], []

    def _load_ocr_adapter(self, adapter_path: str) -> Any:
        module_name, separator, function_name = adapter_path.partition(":")
        if not separator or not module_name or not function_name:
            raise ValueError("MinerU OCR 适配器格式应为 module:function")
        module = importlib.import_module(module_name)
        return getattr(module, function_name)

    def _normalize_ocr_elements(self, raw_result: Any, default_page_no: int) -> list[PdfElement]:
        if raw_result is None:
            return []
        if isinstance(raw_result, str):
            text = self._normalize_text(raw_result)
            return [self._ocr_text_element(default_page_no, text)] if text else []
        if not isinstance(raw_result, list):
            return []

        elements: list[PdfElement] = []
        for item in raw_result:
            if isinstance(item, str):
                text = self._normalize_text(item)
                if text:
                    elements.append(self._ocr_text_element(default_page_no, text))
                continue
            if not isinstance(item, dict):
                continue
            text = self._normalize_text(str(item.get("text") or ""))
            if not text:
                continue
            page_no = int(item.get("page_no") or item.get("pageNo") or default_page_no)
            bbox_value = item.get("bbox") if isinstance(item.get("bbox"), (list, tuple)) else None
            bbox = tuple(float(value) for value in bbox_value[:4]) if bbox_value and len(bbox_value) >= 4 else (0.0, 0.0, 0.0, 0.0)
            element = PdfElement(
                page_no=page_no,
                element_type=str(item.get("element_type") or item.get("type") or "paragraph"),
                text=text,
                bbox=bbox,
                font_size=float(item.get("font_size") or item.get("fontSize") or 0),
                bold=bool(item.get("bold") or False),
                source="mineru",
                metadata={key: value for key, value in item.items() if key not in {"text", "element_type", "type"}},
            )
            if element.element_type not in {"title", "paragraph", "table", "page_number"}:
                element.element_type = "paragraph"
            if element.element_type == "paragraph" and self._looks_like_heading(text):
                element.element_type = "title"
            elements.append(element)
        return elements

    def _ocr_text_element(self, page_no: int, text: str) -> PdfElement:
        element_type = "title" if self._looks_like_heading(text) else "paragraph"
        return PdfElement(
            page_no=page_no,
            element_type=element_type,
            text=text,
            bbox=(0.0, 0.0, 0.0, 0.0),
            source="mineru",
        )

    def _open_pdfplumber(self, file_bytes: bytes) -> Any | None:
        try:
            import pdfplumber

            return pdfplumber.open(io.BytesIO(file_bytes))
        except ImportError:
            logger.warning("pdfplumber 未安装，PDF 使用 PyMuPDF 文本兜底抽取")
            return None
        except Exception as exception:
            logger.warning("pdfplumber 打开 PDF 失败，使用 PyMuPDF 文本兜底抽取 error=%s", exception)
            return None

    def _extract_pdf_table_elements(
        self,
        plumber_page: Any | None,
        fitz_page: Any,
        page_no: int,
    ) -> tuple[list[PdfElement], list[tuple[float, float, float, float]]]:
        if plumber_page is not None:
            try:
                table_elements: list[PdfElement] = []
                table_bboxes: list[tuple[float, float, float, float]] = []
                for table in plumber_page.find_tables():
                    rows = table.extract()
                    table_text = self._table_text(rows[0] if rows else [], rows[1:] if len(rows) > 1 else [])
                    bbox = tuple(table.bbox)
                    if not self._is_meaningful_pdf_table(rows, table_text):
                        continue
                    table_bboxes.append(bbox)
                    table_elements.append(PdfElement(page_no=page_no, element_type="table", text=table_text, bbox=bbox))
                return table_elements, table_bboxes
            except Exception as exception:
                logger.debug("pdfplumber 表格抽取失败 page_no=%s error=%s", page_no, exception)

        try:
            table_elements = []
            table_bboxes = []
            table_finder = fitz_page.find_tables()
            for table in table_finder.tables:
                rows = table.extract()
                table_text = self._table_text(rows[0] if rows else [], rows[1:] if len(rows) > 1 else [])
                bbox = tuple(table.bbox)
                if not self._is_meaningful_pdf_table(rows, table_text):
                    continue
                table_bboxes.append(bbox)
                table_elements.append(PdfElement(page_no=page_no, element_type="table", text=table_text, bbox=bbox))
            return table_elements, table_bboxes
        except Exception as exception:
            logger.debug("PyMuPDF 表格抽取失败 page_no=%s error=%s", page_no, exception)
            return [], []

    def _extract_pdf_text_lines_from_plumber(
        self,
        plumber_page: Any,
        page_no: int,
        table_bboxes: list[tuple[float, float, float, float]],
    ) -> list[PdfElement]:
        background_shapes = getattr(plumber_page, "rects", None) or []
        try:
            chars = plumber_page.dedupe_chars().chars
        except Exception:
            chars = plumber_page.chars
        chars = [char for char in chars if self._has_visible_pdf_char(char, background_shapes)]
        if self._is_garbled_pdf_text("".join(char.get("text", "") for char in chars), threshold=0.35):
            logger.warning("PDF 页面文本疑似乱码，跳过 pdfplumber 字符抽取 page_no=%s", page_no)
            return []

        lines: list[list[dict[str, Any]]] = []
        for char in sorted(chars, key=lambda item: (float(item.get("top", 0)), float(item.get("x0", 0)))):
            bbox = self._plumber_char_bbox(char)
            if self._inside_any_bbox(bbox, table_bboxes):
                continue
            if not lines:
                lines.append([char])
                continue
            current = lines[-1]
            current_top = sum(float(item.get("top", 0)) for item in current) / len(current)
            current_height = max(float(item.get("bottom", 0)) - float(item.get("top", 0)) for item in current)
            tolerance = max(2.0, current_height * 0.55)
            if abs(float(char.get("top", 0)) - current_top) <= tolerance:
                current.append(char)
            else:
                lines.append([char])

        elements: list[PdfElement] = []
        for line_chars in lines:
            line_chars = sorted(line_chars, key=lambda item: float(item.get("x0", 0)))
            text = self._normalize_text(self._join_pdf_chars(line_chars))
            if not text or self._is_garbled_pdf_text(text, threshold=0.5):
                continue
            x0 = min(float(char.get("x0", 0)) for char in line_chars)
            top = min(float(char.get("top", 0)) for char in line_chars)
            x1 = max(float(char.get("x1", 0)) for char in line_chars)
            bottom = max(float(char.get("bottom", 0)) for char in line_chars)
            font_size = max(float(char.get("size", 0)) for char in line_chars)
            bold = any(self._pdf_font_is_bold(str(char.get("fontname", ""))) for char in line_chars)
            bbox = (x0, top, x1, bottom)
            elements.append(
                PdfElement(
                    page_no=page_no,
                    element_type="paragraph",
                    text=text,
                    bbox=bbox,
                    font_size=font_size,
                    bold=bold,
                    metadata={
                        "has_colored_background": self._pdf_bbox_has_colored_background(bbox, background_shapes),
                        "page_width": float(getattr(plumber_page, "width", 0.0)),
                    },
                )
            )
        return self._merge_pdf_inline_lines(elements)

    def _extract_pdf_text_lines_from_fitz(
        self,
        fitz_page: Any,
        page_no: int,
        table_bboxes: list[tuple[float, float, float, float]],
    ) -> list[PdfElement]:
        elements: list[PdfElement] = []
        drawing_shapes = self._pdf_fitz_drawing_shapes(fitz_page)
        for block in fitz_page.get_text("dict").get("blocks", []):
            bbox = tuple(block.get("bbox", (0, 0, 0, 0)))
            if block.get("type") != 0 or self._inside_any_bbox(bbox, table_bboxes):
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                text = self._normalize_text("".join(span.get("text", "") for span in spans))
                if not text or self._is_garbled_pdf_text(text, threshold=0.5):
                    continue
                font_size = max((float(span.get("size", 0)) for span in spans), default=0.0)
                bold = any(
                    self._pdf_font_is_bold(str(span.get("font", ""))) or int(span.get("flags", 0)) & 16
                    for span in spans
                )
                line_bbox = tuple(line.get("bbox", bbox))
                elements.append(
                    PdfElement(
                        page_no=page_no,
                        element_type="paragraph",
                        text=text,
                        bbox=line_bbox,
                        font_size=font_size,
                        bold=bold,
                        metadata={
                            "has_colored_background": self._pdf_bbox_has_colored_background(
                                line_bbox,
                                drawing_shapes,
                            ),
                            "page_width": float(fitz_page.rect.width),
                        },
                    )
                )
        return self._merge_pdf_inline_lines(elements)

    def _extract_pdf_image_elements(self, fitz_page: Any, page_no: int) -> list[PdfElement]:
        image_elements: list[PdfElement] = []
        seen_bboxes: set[tuple[int, int, int, int]] = set()

        for image_index, image in enumerate(fitz_page.get_images(full=True), start=1):
            xref = image[0] if image else None
            if xref is None:
                continue
            for rect in fitz_page.get_image_rects(xref):
                bbox = (float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1))
                bbox_key = self._pdf_bbox_key(bbox)
                if bbox_key in seen_bboxes:
                    continue
                image_info = self._pdf_image_info(image)
                if not self._is_pdf_image_block(image_info, bbox, fitz_page.rect):
                    continue
                seen_bboxes.add(bbox_key)
                image_elements.append(
                    PdfElement(
                        page_no=page_no,
                        element_type="image",
                        text=f"[图片] 第 {page_no} 页图片元素",
                        bbox=bbox,
                        source="native_image",
                        metadata={
                            "image_index": image_index,
                            "xref": xref,
                            "width": image_info.get("width"),
                            "height": image_info.get("height"),
                            "ext": image_info.get("ext"),
                        },
                    )
                )

        for block_index, block in enumerate(fitz_page.get_text("dict").get("blocks", []), start=1):
            bbox = tuple(block.get("bbox", (0, 0, 0, 0)))
            bbox_key = self._pdf_bbox_key(bbox)
            if bbox_key in seen_bboxes:
                continue
            if not self._is_pdf_image_block(block, bbox, fitz_page.rect):
                continue
            seen_bboxes.add(bbox_key)
            image_elements.append(
                PdfElement(
                    page_no=page_no,
                    element_type="image",
                    text=f"[图片] 第 {page_no} 页图片元素",
                    bbox=bbox,
                    source="native_image",
                    metadata={
                        "image_index": len(image_elements) + 1,
                        "block_index": block_index,
                        "width": block.get("width"),
                        "height": block.get("height"),
                        "ext": block.get("ext"),
                    },
                )
            )
        return image_elements

    def _pdf_image_info(self, image: tuple[Any, ...]) -> dict[str, Any]:
        return {
            "xref": image[0] if len(image) > 0 else None,
            "width": image[2] if len(image) > 2 else None,
            "height": image[3] if len(image) > 3 else None,
            "bpc": image[4] if len(image) > 4 else None,
            "colorspace": image[5] if len(image) > 5 else None,
            "alt_colorspace": image[6] if len(image) > 6 else None,
            "name": image[7] if len(image) > 7 else None,
            "filter": image[8] if len(image) > 8 else None,
            "ext": image[9] if len(image) > 9 else None,
        }

    def _pdf_bbox_key(self, bbox: tuple[float, float, float, float]) -> tuple[int, int, int, int]:
        return tuple(round(value) for value in bbox)

    def _merge_pdf_inline_lines(self, elements: list[PdfElement]) -> list[PdfElement]:
        if not elements:
            return []
        merged: list[PdfElement] = []
        for element in sorted(elements, key=lambda item: (item.page_no, item.bbox[1], item.bbox[0])):
            if not merged:
                merged.append(element)
                continue
            previous = merged[-1]
            same_page = previous.page_no == element.page_no
            same_baseline = abs(previous.bbox[1] - element.bbox[1]) <= max(previous.font_size, element.font_size, 8) * 0.45
            close_x = 0 <= element.bbox[0] - previous.bbox[2] <= max(previous.font_size, element.font_size, 8) * 2
            same_style = abs(previous.font_size - element.font_size) <= 1.0 and previous.bold == element.bold
            if same_page and same_baseline and close_x and same_style:
                previous.text = self._normalize_text(f"{previous.text} {element.text}")
                previous.bbox = (
                    min(previous.bbox[0], element.bbox[0]),
                    min(previous.bbox[1], element.bbox[1]),
                    max(previous.bbox[2], element.bbox[2]),
                    max(previous.bbox[3], element.bbox[3]),
                )
                previous.metadata["has_colored_background"] = bool(
                    previous.metadata.get("has_colored_background")
                    or element.metadata.get("has_colored_background")
                )
                continue
            merged.append(element)
        return merged

    def _has_visible_pdf_char(self, char: dict[str, Any], background_shapes: list[Any] | None = None) -> bool:
        text = char.get("text", "")
        if not text:
            return False
        # 白字可能是不可见辅助层，也可能是深色标题条上的真实标题；存在彩色背景时必须保留。
        non_stroking_color = char.get("non_stroking_color")
        stroking_color = char.get("stroking_color")
        if self._is_white_pdf_color(non_stroking_color) and self._is_white_pdf_color(stroking_color):
            return self._pdf_bbox_has_colored_background(self._plumber_char_bbox(char), background_shapes or [])
        return True

    def _pdf_fitz_drawing_shapes(self, fitz_page: Any) -> list[dict[str, Any]]:
        try:
            return [
                {"bbox": tuple(drawing["rect"]), "fill": drawing.get("fill")}
                for drawing in fitz_page.get_drawings()
                if drawing.get("rect") is not None and drawing.get("fill") is not None
            ]
        except Exception as exception:
            logger.debug("PDF 页面背景图形抽取失败 page_no=%s error=%s", fitz_page.number + 1, exception)
            return []

    def _pdf_bbox_has_colored_background(
        self,
        bbox: tuple[float, float, float, float],
        background_shapes: list[Any],
    ) -> bool:
        x0, y0, x1, y1 = bbox
        text_area = max((x1 - x0) * (y1 - y0), 1.0)
        for shape in background_shapes:
            if isinstance(shape, dict) and "bbox" in shape:
                shape_bbox = shape["bbox"]
                fill_color = shape.get("fill")
            elif isinstance(shape, dict):
                shape_bbox = (
                    shape.get("x0", 0.0),
                    shape.get("top", 0.0),
                    shape.get("x1", 0.0),
                    shape.get("bottom", 0.0),
                )
                fill_color = shape.get("non_stroking_color") or shape.get("fill_color")
            else:
                continue
            if not self._is_colored_pdf_fill(fill_color):
                continue
            sx0, sy0, sx1, sy1 = (float(value) for value in shape_bbox)
            overlap_width = max(0.0, min(x1, sx1) - max(x0, sx0))
            overlap_height = max(0.0, min(y1, sy1) - max(y0, sy0))
            if overlap_width * overlap_height / text_area >= 0.5:
                return True
        return False

    def _is_colored_pdf_fill(self, color: Any) -> bool:
        if color is None or self._is_white_pdf_color(color):
            return False
        if isinstance(color, (int, float)):
            return float(color) < 0.95
        if isinstance(color, (list, tuple)):
            numeric = [float(item) for item in color if isinstance(item, (int, float))]
            return bool(numeric) and any(item < 0.95 for item in numeric)
        return False

    def _is_white_pdf_color(self, color: Any) -> bool:
        if color is None:
            return False
        if isinstance(color, (int, float)):
            return float(color) >= 0.98
        if isinstance(color, (list, tuple)) and color:
            numeric = [float(item) for item in color if isinstance(item, (int, float))]
            return bool(numeric) and all(item >= 0.98 for item in numeric)
        return False

    def _plumber_char_bbox(self, char: dict[str, Any]) -> tuple[float, float, float, float]:
        return (
            float(char.get("x0", 0)),
            float(char.get("top", 0)),
            float(char.get("x1", 0)),
            float(char.get("bottom", 0)),
        )

    def _join_pdf_chars(self, chars: list[dict[str, Any]]) -> str:
        text_parts: list[str] = []
        previous: dict[str, Any] | None = None
        for char in chars:
            value = char.get("text", "")
            if not value:
                continue
            if previous is not None and self._needs_space_between_pdf_chars(previous, char):
                text_parts.append(" ")
            text_parts.append(value)
            previous = char
        return "".join(text_parts)

    def _needs_space_between_pdf_chars(self, left: dict[str, Any], right: dict[str, Any]) -> bool:
        left_text = left.get("text", "")
        right_text = right.get("text", "")
        if not left_text or not right_text:
            return False
        if not re.match(r"[A-Za-z0-9,.;:%)]", left_text[-1]):
            return False
        if not re.match(r"[A-Za-z0-9(]", right_text[0]):
            return False
        gap = float(right.get("x0", 0)) - float(left.get("x1", 0))
        char_width = max(float(left.get("width", 0)), float(right.get("width", 0)), 1.0)
        return gap >= char_width * 0.45

    def _pdf_font_is_bold(self, font_name: str) -> bool:
        normalized = font_name.lower()
        return any(marker in normalized for marker in ("bold", "black", "heavy", "semibold", "demi"))

    def _is_garbled_pdf_text(self, text: str, threshold: float) -> bool:
        if not text or not text.strip():
            return False
        if re.search(r"\(cid\s*:\s*\d+\s*\)", text):
            return True
        total = 0
        garbled = 0
        for char in text:
            if char.isspace():
                continue
            total += 1
            if self._is_garbled_pdf_char(char):
                garbled += 1
        return total > 0 and garbled / total >= threshold

    def _is_garbled_pdf_char(self, char: str) -> bool:
        code_point = ord(char)
        if 0xE000 <= code_point <= 0xF8FF:
            return True
        if 0xF0000 <= code_point <= 0x10FFFF:
            return True
        if code_point == 0xFFFD:
            return True
        if code_point < 0x20 and char not in ("\t", "\n", "\r"):
            return True
        if 0x80 <= code_point <= 0x9F:
            return True
        return unicodedata.category(char) in {"Cn", "Cs"}

    def _classify_pdf_text_element(self, element: PdfElement, base_font_size: float) -> PdfElement:
        text = element.text.strip()
        if self._looks_like_page_number(text) or self._looks_like_pdf_noise_text(text):
            element.element_type = "page_number"
            return element
        if self._looks_like_pdf_heading(text, element, base_font_size):
            element.element_type = "title"
            return element
        if self._looks_like_table_line(text):
            element.element_type = "table"
            return element
        element.element_type = "paragraph"
        return element

    def _sections_from_pdf_parse_blocks(self, parse_blocks: list[PdfParseBlock]) -> list[ParsedSection]:
        sections: list[ParsedSection] = []
        title_stack: list[tuple[int, str, str]] = []
        current_section: ParsedSection | None = None
        title_level_map = self._pdf_title_level_map(parse_blocks)
        for parse_block in parse_blocks:
            parse_block.local_directory_tree = self._pdf_parse_block_local_directory_tree(parse_block, title_level_map)

        def create_section(title: str, level: int, parent_section_id: str | None) -> ParsedSection:
            parent_titles = [item[1] for item in title_stack if item[0] < level]
            title_path = " / ".join(parent_titles + [title]) if parent_titles else title
            section = ParsedSection(
                section_no=len(sections) + 1,
                section_id=self._uuid32(),
                parent_section_id=parent_section_id,
                title=title,
                title_path=title_path,
                level=level,
                blocks=[],
            )
            sections.append(section)
            return section

        def ensure_section() -> ParsedSection:
            nonlocal current_section, title_stack
            if current_section is not None:
                return current_section
            current_section = create_section(UNRECOGNIZED_DIRECTORY, 1, None)
            title_stack = [(1, current_section.title, current_section.section_id)]
            return current_section

        def flush_content(
            parse_block: PdfParseBlock,
            section: ParsedSection | None,
            content: list[str],
            page_numbers: list[int],
        ) -> None:
            text = "\n".join(item for item in content if item).strip()
            if not text:
                content.clear()
                page_numbers.clear()
                return
            target_section = section or ensure_section()
            target_section.blocks.append(
                ParsedBlock(
                    block_id=self._uuid32(),
                    block_name=f"{parse_block.block_name} · {target_section.title}",
                    block_type="pdf_parse_block",
                    title_path=target_section.title_path,
                    text=text,
                    parse_block_id=parse_block.block_id,
                    parse_block_name=parse_block.block_name,
                    parse_block_page_start=parse_block.page_start,
                    parse_block_page_end=parse_block.page_end,
                    parse_block_metadata=self._pdf_parse_block_metadata(parse_block),
                    page_start=min(page_numbers) if page_numbers else parse_block.page_start,
                    page_end=max(page_numbers) if page_numbers else parse_block.page_end,
                )
            )
            content.clear()
            page_numbers.clear()

        for parse_block in parse_blocks:
            content: list[str] = []
            page_numbers: list[int] = []
            for element in parse_block.elements:
                if element.element_type == "page_number":
                    continue
                if self._is_pdf_title_candidate(element, self._pdf_parse_block_body_font_size(parse_block)):
                    element.element_type = "title"
                    flush_content(parse_block, current_section, content, page_numbers)
                    title = self._clean_heading(element.text)
                    level = title_level_map.get(
                        self._pdf_title_element_key(element),
                        self._pdf_heading_level(element, title_stack),
                    )
                    # 部分 PDF 只有二级标题标记。将首个标题提升为文档根节点，后续同级标题才能形成真实父子树。
                    if not sections and not title_stack and level > 1:
                        level = 1
                    title_stack = [item for item in title_stack if item[0] < level]
                    parent_section_id = title_stack[-1][2] if title_stack else None
                    current_section = create_section(title, level, parent_section_id)
                    title_stack.append((level, title, current_section.section_id))
                    continue
                element_text = self._pdf_element_text(element)
                if element_text:
                    content.append(element_text)
                    page_numbers.append(element.page_no)
            flush_content(parse_block, current_section, content, page_numbers)

        if not sections:
            return []
        if not any(section.blocks for section in sections):
            first_section = sections[0]
            first_section.blocks.append(
                ParsedBlock(
                    block_id=self._uuid32(),
                    block_name=f"{first_section.title_path} · 标题",
                    block_type="title",
                    title_path=first_section.title_path,
                    text=first_section.title,
                )
            )
        return sections

    def _pdf_element_text(self, element: PdfElement) -> str:
        if element.element_type == "image":
            return ""
        if element.element_type == "table":
            return f"[表格]\n{element.text}"
        if element.element_type == "paragraph":
            return element.text
        return ""

    def _pdf_parse_block_title_summary(self, elements: list[PdfElement]) -> str:
        base_font_size = self._pdf_elements_body_font_size(elements)
        titles = [
            self._clean_heading(item.text)
            for item in elements
            if self._is_pdf_title_candidate(item, base_font_size)
        ]
        unique_titles: list[str] = []
        for title in titles:
            if title not in unique_titles:
                unique_titles.append(title)
            if len(unique_titles) >= 2:
                break
        return "、".join(unique_titles) if unique_titles else UNRECOGNIZED_DIRECTORY

    def _pdf_title_candidates(self, elements: list[PdfElement]) -> list[str]:
        candidates: list[str] = []
        base_font_size = self._pdf_elements_body_font_size(elements)
        for element in elements:
            if not self._is_pdf_title_candidate(element, base_font_size):
                continue
            element.element_type = "title"
            title = self._clean_heading(element.text)
            if title and title not in candidates:
                candidates.append(title)
        return candidates[:20]

    def _pdf_element_counts(self, elements: list[PdfElement]) -> dict[str, int]:
        counts: dict[str, int] = {}
        for element in elements:
            counts[element.element_type] = counts.get(element.element_type, 0) + 1
        return counts

    def _pdf_parse_block_metadata(self, parse_block: PdfParseBlock) -> dict[str, Any]:
        return {
            "parse_block_type": "pdf_page_range",
            "title_candidates": parse_block.title_candidates,
            "element_counts": parse_block.element_counts,
            "ocr_status": parse_block.ocr_status,
            "ocr_message": parse_block.ocr_message,
            "local_directory_tree": parse_block.local_directory_tree,
        }

    def _pdf_parse_block_local_directory_tree(
        self,
        parse_block: PdfParseBlock,
        title_level_map: dict[tuple[Any, ...], int],
    ) -> list[dict[str, Any]]:
        base_font_size = self._pdf_parse_block_body_font_size(parse_block)
        title_elements = [
            item
            for item in parse_block.elements
            if self._is_pdf_title_candidate(item, base_font_size)
        ]
        if not title_elements:
            return []

        roots: list[dict[str, Any]] = []
        stack: list[dict[str, Any]] = []
        for element in title_elements:
            level = title_level_map.get(self._pdf_title_element_key(element), self._pdf_heading_level(element, []))
            node = {
                "title": self._clean_heading(element.text),
                "level": level,
                "page_start": element.page_no,
                "page_end": element.page_no,
                "source": element.source,
                "children": [],
            }
            while stack and int(stack[-1]["level"]) >= level:
                stack.pop()
            if stack:
                stack[-1]["children"].append(node)
            else:
                roots.append(node)
            stack.append(node)
        return roots

    def _pdf_title_level_map(self, parse_blocks: list[PdfParseBlock]) -> dict[tuple[Any, ...], int]:
        all_elements = [element for parse_block in parse_blocks for element in parse_block.elements]
        base_font_size = self._pdf_elements_body_font_size(all_elements)
        title_elements = [
            element
            for parse_block in parse_blocks
            for element in parse_block.elements
            if self._is_pdf_title_candidate(element, base_font_size)
        ]
        if not title_elements:
            return {}

        # MinerU、Docling 等版面解析器能直接给出标题层级时，优先保留解析器判断；
        # 其余原生 PDF 标题再结合字号、粗体和位置推断，避免外部解析结果被样式启发式覆盖。
        level_map: dict[tuple[Any, ...], int] = {}
        inferred_title_elements: list[PdfElement] = []
        for element in title_elements:
            heading_level = element.metadata.get("heading_level")
            if heading_level is None:
                inferred_title_elements.append(element)
                continue
            try:
                level_map[self._pdf_title_element_key(element)] = max(1, min(int(heading_level), 6))
            except (TypeError, ValueError):
                inferred_title_elements.append(element)

        if not inferred_title_elements:
            return level_map

        max_font_size = max((item.font_size for item in inferred_title_elements), default=0)
        document_title_key = self._pdf_document_title_key(inferred_title_elements, max_font_size)
        document_title_offset = 1 if document_title_key else 0
        style_rank = self._pdf_title_style_rank(inferred_title_elements)
        for element in inferred_title_elements:
            key = self._pdf_title_element_key(element)
            if key == document_title_key:
                level_map[key] = 1
                continue
            numbered_level = self._pdf_numbered_heading_level(element.text)
            if numbered_level is not None:
                level_map[key] = max(1, min(numbered_level + document_title_offset, 6))
                continue
            level_map[key] = max(1, min(style_rank.get(self._pdf_title_style_key(element), 2) + document_title_offset, 6))
        return level_map

    def _pdf_document_title_key(self, title_elements: list[PdfElement], max_font_size: float) -> tuple[Any, ...] | None:
        candidates = [
            element
            for element in title_elements
            if element.page_no == 1
            and element.bbox[1] <= 160
            and element.font_size >= max_font_size - 1
            and self._pdf_numbered_heading_level(element.text) is None
            and len(self._clean_heading(element.text)) >= 4
        ]
        if not candidates:
            return None
        title = sorted(candidates, key=lambda item: (item.bbox[1], item.bbox[0]))[0]
        return self._pdf_title_element_key(title)

    def _pdf_title_element_key(self, element: PdfElement) -> tuple[Any, ...]:
        return (
            element.page_no,
            self._clean_heading(element.text),
            round(element.bbox[0], 1),
            round(element.bbox[1], 1),
            round(element.bbox[2], 1),
            round(element.bbox[3], 1),
        )

    def _pdf_title_style_key(self, element: PdfElement) -> tuple[Any, ...]:
        width = element.bbox[2] - element.bbox[0]
        height = element.bbox[3] - element.bbox[1]
        return (
            round(element.font_size * 2) / 2,
            element.bold,
            round(element.bbox[0] / 12),
            round(width / 40),
            round(height / 4),
        )

    def _pdf_title_style_rank(self, title_elements: list[PdfElement]) -> dict[tuple[Any, ...], int]:
        style_items: dict[tuple[Any, ...], dict[str, Any]] = {}
        for element in title_elements:
            style_key = self._pdf_title_style_key(element)
            current = style_items.setdefault(
                style_key,
                {
                    "font_size": element.font_size,
                    "bold": element.bold,
                    "min_y": element.bbox[1],
                    "count": 0,
                },
            )
            current["font_size"] = max(float(current["font_size"]), element.font_size)
            current["bold"] = bool(current["bold"]) or element.bold
            current["min_y"] = min(float(current["min_y"]), element.bbox[1])
            current["count"] = int(current["count"]) + 1
        ordered_styles = sorted(
            style_items.items(),
            key=lambda item: (-float(item[1]["font_size"]), not bool(item[1]["bold"]), float(item[1]["min_y"])),
        )
        return {style_key: min(index + 1, 6) for index, (style_key, _) in enumerate(ordered_styles)}

    def _pdf_numbered_heading_level(self, text: str) -> int | None:
        cleaned = self._clean_heading(text)
        if re.match(r"^第[一二三四五六七八九十百千万\d]+[章节篇部]", cleaned):
            return 1
        if re.match(r"^[一二三四五六七八九十]+[、.．]\s*", cleaned):
            return 1
        numbered_match = re.match(r"^(\d+(?:[.．]\d+)*)(?:[.．、]|\s+)", cleaned)
        if not numbered_match:
            return None
        number_text = numbered_match.group(1)
        return max(1, min(number_text.count(".") + number_text.count("．") + 1, 6))

    def _pdf_heading_level(self, element: PdfElement, title_stack: list[tuple[int, str, str]]) -> int:
        heading_level = element.metadata.get("heading_level")
        if heading_level is not None:
            try:
                return max(1, min(int(heading_level), 6))
            except (TypeError, ValueError):
                pass
        numbered_level = self._pdf_numbered_heading_level(element.text)
        if numbered_level is not None:
            return numbered_level
        if element.font_size >= 18:
            return 1
        if element.font_size >= 14:
            return min(2, len(title_stack) + 1)
        return min(3, len(title_stack) + 1)

    def _looks_like_pdf_heading(self, text: str, element: PdfElement, base_font_size: float) -> bool:
        if not text or len(text) > 60:
            return False
        if self._looks_like_page_number(text) or self._looks_like_pdf_noise_text(text):
            return False
        if self._looks_like_pdf_list_item(text):
            return False
        if self._looks_like_heading(text):
            return True
        # 简历、表单中的“电话：xxx”“项目简介：xxx”等键值字段即使居中或粗体，也不是目录标题。
        if re.match(r"^.{1,12}[：:]\s*\S+", text):
            return False
        if element.metadata.get("has_colored_background") and len(text) <= 40:
            return True
        if base_font_size > 0 and element.font_size >= base_font_size + 3 and len(text) <= 40:
            return True
        page_width = float(element.metadata.get("page_width") or 0.0)
        element_center = (element.bbox[0] + element.bbox[2]) / 2
        is_centered = page_width > 0 and abs(element_center - page_width / 2) <= page_width * 0.08
        if is_centered and element.bold and len(text) <= 40:
            return True
        if re.fullmatch(r"[\u4e00-\u9fffA-Za-z0-9 /（）()_-]{2,24}", text) and (
            (base_font_size > 0 and element.font_size >= base_font_size + 1.5) or element.bold
        ):
            return True
        return False

    def _is_pdf_title_candidate(self, element: PdfElement, base_font_size: float) -> bool:
        if element.element_type in {"image", "table", "page_number"}:
            return False
        text = self._normalize_text(element.text)
        if not text or self._looks_like_page_number(text) or self._looks_like_pdf_noise_text(text):
            return False
        if self._looks_like_pdf_list_item(text):
            return False
        if element.metadata.get("heading_level") is not None:
            return True
        if self._looks_like_pdf_heading(text, element, base_font_size):
            return True
        # 上游 title 仅作为弱证据：外部版面解析器的标题标签还需短文本约束；
        # 原生 PDF 则必须再有粗体或相对大字号支持，不能只信 element_type。
        if element.element_type != "title" or len(text) > 40:
            return False
        if element.source != "native":
            return True
        return element.bold or (base_font_size > 0 and element.font_size >= base_font_size + 1.0)

    def _pdf_parse_block_body_font_size(self, parse_block: PdfParseBlock) -> float:
        return self._pdf_elements_body_font_size(parse_block.elements)

    def _pdf_elements_body_font_size(self, elements: list[PdfElement]) -> float:
        body_sizes = [
            element.font_size
            for element in elements
            if element.element_type == "paragraph"
            and element.font_size > 0
            and not element.metadata.get("heading_level")
            and not element.metadata.get("has_colored_background")
        ]
        return self._pdf_body_font_size(body_sizes)

    def _looks_like_pdf_list_item(self, text: str) -> bool:
        stripped = text.strip()
        if not re.match(r"^\d+[.．]\s*", stripped):
            return False
        if len(stripped) > 24:
            return True
        return bool(re.search(r"[，,。；;：:].{4,}", stripped))

    def _pdf_body_font_size(self, font_sizes: list[float]) -> float:
        if not font_sizes:
            return 0.0
        sorted_sizes = sorted(font_sizes)
        middle = len(sorted_sizes) // 2
        if len(sorted_sizes) % 2:
            return sorted_sizes[middle]
        return (sorted_sizes[middle - 1] + sorted_sizes[middle]) / 2

    def _is_meaningful_pdf_table(self, rows: list[list[Any]], table_text: str) -> bool:
        if not table_text:
            return False
        non_empty_rows = [row for row in rows if any(self._cell_text(cell) for cell in row)]
        if len(non_empty_rows) < 2:
            return False
        max_columns = max((len([cell for cell in row if self._cell_text(cell)]) for row in non_empty_rows), default=0)
        if max_columns < 2:
            return False
        return True

    def _is_pdf_image_block(
        self,
        block: dict[str, Any],
        bbox: tuple[float, float, float, float],
        page_rect: Any | None = None,
    ) -> bool:
        has_image_payload = any(block.get(key) is not None for key in ("image", "xref", "ext", "width", "height", "size"))
        if not has_image_payload:
            return False
        x0, y0, x1, y1 = bbox
        width = x1 - x0
        height = y1 - y0
        if width < 32 or height < 32:
            return False
        if page_rect is not None:
            page_width = float(page_rect.width)
            page_height = float(page_rect.height)
            # 简历、表单里常见的彩色标题背景条会被 PyMuPDF 标成图片块，
            # 这类装饰块不能作为正文图片进入分片或目录结构。
            if width >= page_width * 0.45 and height <= page_height * 0.08:
                return False
            if height >= page_height * 0.45 and width <= page_width * 0.08:
                return False
        aspect_ratio = width / max(height, 1)
        if aspect_ratio >= 6 or aspect_ratio <= 0.16:
            return False
        return True

    def _looks_like_pdf_noise_text(self, text: str) -> bool:
        stripped = text.strip()
        if len(stripped) <= 1:
            return True
        if re.fullmatch(r"[\W_]+", stripped):
            return True
        if re.fullmatch(r"(page|p)\s*\d+", stripped, re.IGNORECASE):
            return True
        return False

    def _looks_like_page_number(self, text: str) -> bool:
        return bool(re.fullmatch(r"(\d+|第\s*\d+\s*页|-\s*\d+\s*-)", text.strip()))

    def _looks_like_table_line(self, text: str) -> bool:
        if "|" in text or "\t" in text:
            return True
        if len(re.findall(r"\s{2,}", text)) >= 2:
            return True
        return bool(re.search(r"[:：]\s*[^:：]+[:：]", text))

    def _inside_any_bbox(
        self,
        bbox: tuple[float, float, float, float],
        bboxes: list[tuple[float, float, float, float]],
    ) -> bool:
        x0, y0, x1, y1 = bbox
        for tx0, ty0, tx1, ty1 in bboxes:
            if x0 >= tx0 - 2 and y0 >= ty0 - 2 and x1 <= tx1 + 2 and y1 <= ty1 + 2:
                return True
        return False

    def _split_chunks(self, sections: list[ParsedSection], message: DocumentProcessMessage) -> list[Chunk]:
        max_chars = self._chunk_max_chars(message)
        overlap = min(settings.chunk_overlap_chars, max_chars // 4)
        chunks: list[Chunk] = []
        for section in sections:
            for block in section.blocks:
                text = block.text.strip()
                if not text:
                    continue
                start = 0
                while start < len(text):
                    end = min(len(text), start + max_chars)
                    chunk_text = text[start:end].strip()
                    if chunk_text:
                        chunk_no = len(chunks) + 1
                        chunks.append(
                            Chunk(
                                chunk_id=self._chunk_id(message, chunk_no),
                                chunk_no=chunk_no,
                                chunk_text=chunk_text,
                                section_id=section.section_id,
                                title_path=block.title_path,
                                block_ids=[block.block_id],
                                block_names=[block.block_name or UNASSIGNED_BLOCK],
                                parse_block_id=block.parse_block_id,
                                parse_block_name=block.parse_block_name,
                                token_count=self._estimate_tokens(chunk_text),
                                char_count=len(chunk_text),
                                page_start=block.page_start,
                                page_end=block.page_end,
                                sheet_name=block.sheet_name,
                                row_start=block.row_start,
                                row_end=block.row_end,
                            )
                        )
                    if end >= len(text):
                        break
                    start = max(end - overlap, start + 1)
        return chunks or [
            Chunk(
                chunk_id=self._chunk_id(message, 1),
                chunk_no=1,
                chunk_text=self._fallback_text(message),
                section_id=None,
                title_path=UNRECOGNIZED_STRUCTURE,
                block_ids=[self._uuid32()],
                block_names=[UNASSIGNED_BLOCK],
                parse_block_id=None,
                parse_block_name=None,
                token_count=1,
                char_count=len(self._fallback_text(message)),
            )
        ]

    def _ensure_parse_block_metadata(
        self,
        sections: list[ParsedSection],
        message: DocumentProcessMessage,
    ) -> list[ParsedSection]:
        file_type = message.fileType.lower()
        for index, section in enumerate(sections, start=1):
            if not section.blocks:
                continue
            existing_block_id = next((block.parse_block_id for block in section.blocks if block.parse_block_id), None)
            parse_block_id = existing_block_id or self._uuid32()
            parse_block_name = self._parse_block_name(file_type, index, section)
            title_candidates = [section.title] if section.title and section.title not in {UNRECOGNIZED_STRUCTURE, UNRECOGNIZED_DIRECTORY} else []
            page_starts = [block.page_start for block in section.blocks if block.page_start is not None]
            page_ends = [block.page_end for block in section.blocks if block.page_end is not None]
            row_starts = [block.row_start for block in section.blocks if block.row_start is not None]
            row_ends = [block.row_end for block in section.blocks if block.row_end is not None]
            metadata = {
                "parse_block_type": self._parse_block_type(file_type),
                "title_candidates": title_candidates,
                "element_counts": self._section_block_type_counts(section),
                "ocr_status": "not_required" if file_type != "pdf" else "not_required",
            }
            for block in section.blocks:
                if not block.parse_block_id:
                    block.parse_block_id = parse_block_id
                    block.parse_block_name = parse_block_name
                if block.parse_block_page_start is None and page_starts:
                    block.parse_block_page_start = min(page_starts)
                if block.parse_block_page_end is None and page_ends:
                    block.parse_block_page_end = max(page_ends)
                if not block.parse_block_metadata:
                    block.parse_block_metadata = dict(metadata)
                if row_starts and row_ends:
                    block.parse_block_metadata.setdefault("row_start", min(row_starts))
                    block.parse_block_metadata.setdefault("row_end", max(row_ends))
        return sections

    def _parse_block_name(self, file_type: str, index: int, section: ParsedSection) -> str:
        if file_type in {"docx", "doc", "md", "markdown"}:
            return f"解析块 {index:02d}｜{section.title_path}"
        if file_type in {"xlsx", "xls", "csv"}:
            sheet_name = section.blocks[0].sheet_name if section.blocks else None
            return f"解析块 {index:02d}｜{sheet_name or section.title_path}"
        if file_type == "txt":
            return f"解析块 {index:02d}｜文本段落"
        return f"解析块 {index:02d}｜{section.title_path or UNRECOGNIZED_STRUCTURE}"

    def _parse_block_type(self, file_type: str) -> str:
        if file_type == "pdf":
            return "pdf_page_range"
        if file_type in {"docx", "doc", "md", "markdown"}:
            return "document_heading"
        if file_type in {"xlsx", "xls", "csv"}:
            return "table_row_range"
        return "text_section"

    def _section_block_type_counts(self, section: ParsedSection) -> dict[str, int]:
        counts: dict[str, int] = {}
        for block in section.blocks:
            counts[block.block_type] = counts.get(block.block_type, 0) + 1
        return counts

    def _section_from_text(
        self,
        section_no: int,
        title_path: str,
        text: str,
        page_start: int | None = None,
        page_end: int | None = None,
    ) -> ParsedSection:
        title_path = title_path or UNRECOGNIZED_STRUCTURE
        block = ParsedBlock(
            block_id=self._uuid32(),
            block_name=f"{title_path} · 正文",
            block_type="text",
            title_path=title_path,
            text=text,
            page_start=page_start,
            page_end=page_end,
        )
        return ParsedSection(
            section_no=section_no,
            section_id=self._uuid32(),
            parent_section_id=None,
            title=title_path.split(" / ")[-1],
            title_path=title_path,
            level=max(1, title_path.count(" / ") + 1),
            blocks=[block],
        )

    def _table_section(
        self,
        section_no: int,
        title_path: str,
        sheet_name: str,
        header: list[Any],
        rows: list[list[Any]],
        first_data_row: int,
    ) -> ParsedSection:
        blocks: list[ParsedBlock] = []
        max_rows = max(1, settings.table_chunk_rows)
        for start in range(0, len(rows), max_rows):
            row_slice = rows[start : start + max_rows]
            row_start = first_data_row + start
            row_end = row_start + len(row_slice) - 1
            block_no = len(blocks) + 1
            blocks.append(
                ParsedBlock(
                    block_id=self._uuid32(),
                    block_name=f"{title_path} · 第 {row_start}-{row_end} 行",
                    block_type="table",
                    title_path=title_path,
                    text=self._table_text(header, row_slice),
                    sheet_name=sheet_name,
                    row_start=row_start,
                    row_end=row_end,
                )
            )
        if not blocks:
            blocks.append(
                ParsedBlock(
                    block_id=self._uuid32(),
                    block_name=f"{title_path} · 表头",
                    block_type="table",
                    title_path=title_path,
                    text=self._table_text(header, []),
                    sheet_name=sheet_name,
                    row_start=1,
                    row_end=1,
                )
            )
        return ParsedSection(section_no, self._uuid32(), None, title_path, title_path, 1, blocks)

    def _build_sections_jsonl(self, sections: list[ParsedSection]) -> str:
        rows = []
        for section in sections:
            first_block = section.blocks[0] if section.blocks else None
            rows.append(
                json.dumps(
                    {
                        "section_no": section.section_no,
                        "section_id": section.section_id,
                        "parent_section_id": section.parent_section_id,
                        "title": section.title,
                        "title_path": section.title_path,
                        "level": section.level,
                        "parse_block_id": first_block.parse_block_id if first_block else None,
                        "parse_block_name": first_block.parse_block_name if first_block else None,
                        "page_start": first_block.parse_block_page_start if first_block else None,
                        "page_end": first_block.parse_block_page_end if first_block else None,
                        "block_count": len(section.blocks),
                        "text_preview": "\n".join(block.text for block in section.blocks)[:500],
                    },
                    ensure_ascii=False,
                )
            )
        return "\n".join(rows) + "\n"

    def _build_structure_json(
        self, message: DocumentProcessMessage, sections: list[ParsedSection], parser_name: str
    ) -> str:
        structure = {
            "document_id": message.documentId,
            "version_id": message.versionId,
            "process_id": message.taskId,
            "parser_name": parser_name,
            "directory_tree": self._build_directory_tree(sections),
            "parse_blocks": self._build_parse_block_summary(sections),
            "sections": [
                {
                    "section_no": section.section_no,
                    "section_id": section.section_id,
                    "parent_section_id": section.parent_section_id,
                    "title": section.title,
                    "title_path": section.title_path,
                    "level": section.level,
                    "blocks": [
                        {
                            "block_id": block.block_id,
                            "block_name": block.block_name,
                            "block_type": block.block_type,
                            "parse_block_id": block.parse_block_id,
                            "parse_block_name": block.parse_block_name,
                            "parse_block_page_start": block.parse_block_page_start,
                            "parse_block_page_end": block.parse_block_page_end,
                            "parse_block_metadata": block.parse_block_metadata,
                            "title_path": block.title_path,
                            "page_start": block.page_start,
                            "page_end": block.page_end,
                            "sheet_name": block.sheet_name,
                            "row_start": block.row_start,
                            "row_end": block.row_end,
                            "text_preview": block.text[:500],
                        }
                        for block in section.blocks
                    ],
                }
                for section in sections
            ],
        }
        return json.dumps(structure, ensure_ascii=False, indent=2)

    def _build_directory_tree(self, sections: list[ParsedSection]) -> list[dict[str, Any]]:
        node_map: dict[str, dict[str, Any]] = {}
        roots: list[dict[str, Any]] = []
        for section in sections:
            page_starts = [block.page_start for block in section.blocks if block.page_start is not None]
            page_ends = [block.page_end for block in section.blocks if block.page_end is not None]
            node = {
                "section_no": section.section_no,
                "section_id": section.section_id,
                "parent_section_id": section.parent_section_id,
                "title": section.title,
                "title_path": section.title_path,
                "level": section.level,
                "page_start": min(page_starts) if page_starts else None,
                "page_end": max(page_ends) if page_ends else None,
                "block_count": len(section.blocks),
                "children": [],
            }
            node_map[section.section_id] = node

        for section in sections:
            node = node_map[section.section_id]
            if section.parent_section_id and section.parent_section_id in node_map:
                node_map[section.parent_section_id]["children"].append(node)
            else:
                roots.append(node)
        return roots

    def _build_parse_block_summary(self, sections: list[ParsedSection]) -> list[dict[str, Any]]:
        summary_map: dict[str, dict[str, Any]] = {}
        for section in sections:
            for block in section.blocks:
                if not block.parse_block_id:
                    continue
                current = summary_map.setdefault(
                    block.parse_block_id,
                    {
                        "parse_block_id": block.parse_block_id,
                        "parse_block_name": block.parse_block_name,
                        "page_start": block.parse_block_page_start,
                        "page_end": block.parse_block_page_end,
                        "section_ids": [],
                        "section_titles": [],
                        "metadata": block.parse_block_metadata,
                        "text_preview": "",
                    },
                )
                if section.section_id not in current["section_ids"]:
                    current["section_ids"].append(section.section_id)
                    current["section_titles"].append(section.title)
                if not current["text_preview"] and block.text:
                    current["text_preview"] = block.text[:500]
                if not current.get("metadata") and block.parse_block_metadata:
                    current["metadata"] = block.parse_block_metadata
        return list(summary_map.values())

    def _build_chunks_jsonl(self, message: DocumentProcessMessage, chunks: list[Chunk]) -> str:
        rows = []
        for chunk in chunks:
            rows.append(
                json.dumps(
                    {
                        "chunk_id": chunk.chunk_id,
                        "document_id": message.documentId,
                        "version_id": message.versionId,
                        "process_id": message.taskId,
                        "chunk_no": chunk.chunk_no,
                        "chunk_text": chunk.chunk_text,
                        "section_id": chunk.section_id,
                        "title_path": chunk.title_path,
                        "source_type": message.fileType,
                        "block_ids": chunk.block_ids,
                        "block_names": chunk.block_names,
                        "parse_block_id": chunk.parse_block_id,
                        "parse_block_name": chunk.parse_block_name,
                        "token_count": chunk.token_count,
                        "char_count": chunk.char_count,
                        "page_start": chunk.page_start,
                        "page_end": chunk.page_end,
                        "sheet_name": chunk.sheet_name,
                        "row_start": chunk.row_start,
                        "row_end": chunk.row_end,
                    },
                    ensure_ascii=False,
                )
            )
        return "\n".join(rows) + "\n"

    def _chunk_max_chars(self, message: DocumentProcessMessage) -> int:
        max_chars = settings.chunk_max_chars
        if message.chunkConfigSnapshot:
            try:
                snapshot = json.loads(message.chunkConfigSnapshot)
                max_chars = int(snapshot.get("chunkMaxChars") or snapshot.get("maxChars") or max_chars)
            except Exception:
                logger.warning("分片配置快照解析失败 task_id=%s", message.taskId)
        return max(200, min(max_chars, 4000))

    def _chunk_id(self, message: DocumentProcessMessage, chunk_no: int) -> str:
        return self._uuid32()

    def _current_title_path(self, heading_stack: list[tuple[int, str]]) -> str:
        return " / ".join(title for _, title in heading_stack) if heading_stack else UNRECOGNIZED_STRUCTURE

    def _looks_like_heading(self, line: str) -> bool:
        if not line or len(line) > 80:
            return False
        return bool(
            re.match(
                r"^([一二三四五六七八九十]+[、.．]|第[一二三四五六七八九十\d]+[章节条]|[0-9]+(?:[.．][0-9]+)*[.．、]\s+).+",
                line,
            )
        )

    def _docx_heading_level(self, style_name: str, text: str) -> int | None:
        match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", style_name.lower(), re.IGNORECASE)
        if match:
            level_text = match.group(1) or match.group(2)
            return max(1, min(int(level_text), 6))
        if self._looks_like_heading(text):
            return 1
        return None

    def _is_markdown_table_line(self, line: str) -> bool:
        stripped = line.strip()
        if not stripped:
            return False
        if stripped.startswith("|") and stripped.endswith("|"):
            return True
        return bool(re.match(r"^\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+$", stripped))

    def _clean_heading(self, heading: str) -> str:
        value = re.sub(r"^[#\s]+", "", heading).strip()
        value = re.sub(r"\s+", " ", value)
        return value[:80] or UNRECOGNIZED_STRUCTURE

    def _normalize_text(self, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _decode_text(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk", "big5"):
            try:
                return self._fix_mojibake(file_bytes.decode(encoding).strip())
            except UnicodeDecodeError:
                continue
        return self._fix_mojibake(file_bytes.decode("utf-8", errors="ignore").strip())

    def _fix_mojibake(self, text: str) -> str:
        if not text:
            return text
        markers = ("Ã", "Â", "æ", "è", "å", "ç", "¤", "\ufffd")
        if not any(marker in text for marker in markers):
            return text
        for source_encoding in ("latin1", "cp1252"):
            try:
                recovered = text.encode(source_encoding, errors="ignore").decode("utf-8")
            except UnicodeError:
                continue
            if recovered and recovered.count("\ufffd") < text.count("\ufffd") and self._has_cjk(recovered):
                return recovered
        return text

    def _has_cjk(self, text: str) -> bool:
        return any("\u4e00" <= char <= "\u9fff" for char in text)

    def _uuid32(self) -> str:
        return uuid.uuid4().hex

    def _table_text(self, header: list[Any], rows: list[list[Any]]) -> str:
        header_text = " | ".join(self._cell_text(item) for item in header)
        lines = [f"表头：{header_text}"] if header_text else []
        for row in rows:
            lines.append(" | ".join(self._cell_text(item) for item in row))
        return "\n".join(lines).strip()

    def _cell_text(self, value: Any) -> str:
        return "" if value is None else str(value).strip()

    def _estimate_tokens(self, text: str) -> int:
        ascii_words = re.findall(r"[A-Za-z0-9_]+", text)
        chinese_chars = re.findall(r"[\u4e00-\u9fff]", text)
        return max(1, len(ascii_words) + len(chinese_chars))

    def _fallback_text(self, message: DocumentProcessMessage) -> str:
        return f"{message.fileName} 暂未解析出正文内容。"
